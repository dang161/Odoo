# -*- coding: utf-8 -*-

from odoo import models, fields, api
from odoo.exceptions import UserError, ValidationError
from tempfile import TemporaryFile, NamedTemporaryFile
from docx import Document
# from docxtpl import DocxTemplate, InlineImage, CheckBox, CheckedBox, Tick
from io import BytesIO, StringIO
import logging
from . import intern_utils
from docx.shared import Mm, Inches
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
import os

_logger = logging.getLogger(__name__)
import codecs


class Invoice(models.Model):
    _name = 'intern.invoice'
    _rec_name = 'name'

    custom_id = fields.Char('Mã tự động')

    @api.multi
    def start_translate_form(self):
        view_id = self.env.ref('hh_intern.view_doc_generate_clone').id
        context = self._context.copy()
        return {
            'name': 'form_name',
            'view_type': 'form',
            'view_mode': 'form',
            'res_model': 'intern.invoice',
            'view_id': view_id,
            'type': 'ir.actions.act_window',
            'res_id': self.id,
            'context': context,
        }

    @api.multi
    def create_doc_new(self):
        if self.interns_exam_doc is None or len(self.interns_exam_doc) is 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách thi tuyển")

        _logger.info("CREATE DOC ENDDD")
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document_new?model=intern.invoice&id=%s&filename=%s.zip' % (
                str(self.id), self.name),
            'target': 'self', }

    @api.multi
    def create_doc_new_man(self):

        if self.interns_exam_doc is None or len(self.interns_exam_doc) is 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách thi tuyển")

        _logger.info("CREATE DOC ENDDD")
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document_new?model=intern.invoice&id=%s&filename=%s.zip&gender=nam' % (
                str(self.id), self.name),
            'target': 'self', }

    @api.multi
    def create_doc_new_women(self):

        if self.interns_exam_doc is None or len(self.interns_exam_doc) is 0:
            raise ValidationError("Không có thực tập sinh nào trong danh sách thi tuyển")

        _logger.info("CREATE DOC ENDDD")
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document_new?model=intern.invoice&id=%s&filename=%s.zip&gender=nu' % (
                str(self.id), self.name),
            'target': 'self', }

    @api.model
    def create(self, vals):
        _logger.info("CREATE %s" % str(vals))
        result = super(Invoice, self).create(vals)
        result['custom_id'] = 'ERP-%d' % result.id
        if 'hoso_created' in vals and vals['hoso_created']:
            result['status'] = 2
        return result

    @api.multi
    def write(self, vals):
        pre_promoted = []
        if 'interns_clone' in vals:
            for intern in self.interns_clone:
                if intern.promoted:
                    pre_promoted.append(intern.id)
        tmp = super(Invoice, self).write(vals)
        if 'interns_pass_doc' in vals or 'interns_pass_new' in vals:
            enterprises = []
            for intern in self.interns_pass_doc:
                if intern.enterprise and intern.enterprise.id not in enterprises:
                    enterprises.append(intern.enterprise.id)
            listtmp = sorted(self.interns_pass_doc, key=lambda x: x.sequence_pass)
            counter = 0
            for id in enterprises:
                for i, intern in enumerate(listtmp):
                    if not intern.enterprise:
                        intern.sequence_pass = len(self.interns_pass_doc)
                    elif intern.enterprise.id == id:
                        intern.sequence_pass = counter
                        counter += 1
        if 'interns_clone' in vals:
            last_promoted = []
            for intern in self.interns_clone:
                if self.enterprise_doc and not intern.enterprise:
                    intern.enterprise = self.enterprise_doc
                if intern.promoted:
                    last_promoted.append(intern.id)

            sub_promoted = [item for item in pre_promoted if item not in last_promoted]
            current = []
            if self.promotion_removed:
                test = self.promotion_removed.split(',')
                for x in test:
                    current.append((x))

            current = current + sub_promoted
            lastest = [x for x in current if x not in last_promoted]
            self.promotion_removed = ','.join(map(str, lastest))
        return tmp

    custom_id_2 = fields.Char(string="Mã đơn hàng")  # hh_142
    room_pttt = fields.Many2one(comodel_name="phongban.phongban", string="Phòng PTTT", required=False, )  # hh_143

    name = fields.Char(string="Tên đơn hàng", required=True)  # hh_144
    employee_pttt = fields.Many2one(comodel_name="nhanvien.nhanvien", string="Cán bộ PTTT", required=False, )  # hh_145
    room_td_care = fields.Many2one(comodel_name="phongban.phongban", string="Phòng TD hỗ trợ",
                                   required=False, )  # hh_146

    place_to_work = fields.Char(string="Địa điểm làm việc")  # hh_147

    job_predefine = fields.Many2one(comodel_name="intern.job", string="Ngành nghề xin thư tiến cử",
                                    required=False, )  # hh_148
    job_en = fields.Char("Ngành nghề xin thư tiến cử lấy từ hợp đồng lương (Tiếng Anh)")  # hh_195
    job_jp = fields.Char("Ngành nghề xin thư tiến cử lấy từ hợp đồng lương (Tiếng Nhật)")  # hh_196
    job_vi = fields.Char("Ngành nghề xin thư tiến cử lấy từ hợp đồng lương (Tiếng Việt)")  # hh_197

    @api.onchange('job_predefine')
    def job_change(self):
        if self.job_predefine:
            self.job_en = self.job_predefine.name_en
            self.job_jp = self.job_predefine.name_jp
            self.job_vi = self.job_predefine.name

    day_exam = fields.Char(string="Ngày")  # hh_150
    month_exam = fields.Selection([('01', '01'), ('02', '02'), ('03', '03'), ('04', '04'),
                                   ('05', '05'), ('06', '06'), ('07', '07'), ('08', '08'),
                                   ('09', '09'), ('10', '10'), ('11', '11'), ('12', '12'), ], "Tháng")  # hh_151
    year_exam = fields.Char(string="Năm", required=False, )  # hh_152
    date_confirm_form = fields.Date('Ngày chốt form')  # hh_153
    date_confirm_form1 = fields.Char('Ngày chốt form', store=True, compute='_date_confirm_form1')

    @api.one
    @api.depends('date_confirm_form')
    def _date_confirm_form1(self):
        self.date_confirm_form1 = self.date_confirm_form

    date_exam = fields.Char(string="Ngày thi", store=True, compute='_date_exam')  # hh_149

    @api.one
    @api.depends('day_exam', 'month_exam', 'year_exam')
    def _date_exam(self):
        if self.day_exam and self.month_exam and self.year_exam:
            self.date_exam = u"Ngày %s tháng %s năm %s" % (
                self.day_exam, self.month_exam, self.year_exam)
        elif self.month_exam and self.year_exam:
            self.date_exam = u"Tháng %s năm %s" % (
                self.month_exam, self.year_exam)
        elif self.year_exam:
            self.date_exam = u'Năm %s' % self.year_exam
        else:
            self.date_exam = ""

    date_exam_short = fields.Char("Ngày thi tuyển", store=True, compute='_date_exam_short')

    @api.multi
    @api.depends('day_exam', 'month_exam', 'year_exam')
    def _date_exam_short(self):
        strdate = ""
        for rec in self:
            if rec.day_exam and rec.month_exam and rec.year_exam:
                strdate = datetime.strptime('%s-%s-%s' % (rec.year_exam, rec.month_exam, rec.day_exam),
                                            '%Y-%m-%d')
            else:
                strdate = None
        arr = str(strdate).split()
        self.date_exam_short = arr[0]
        print(self.date_exam_short)

    year_expire = fields.Integer(string="Thời hạn hợp đồng (năm)", required=False, )  # hh_155
    salary_base = fields.Char(string="Lương cơ bản", required=False, )  # hh_156
    salary_real = fields.Char(string="Lương thực lĩnh", required=False, )  # hh_157
    number_man = fields.Integer(string="Số lượng nam", required=False, )  # hh_158
    number_women = fields.Integer(string="Số lượng nữ")  # hh_159
    number_total = fields.Integer(string="Số lượng trúng tuyển", required=False, )  # hh_160
    age_from = fields.Integer(string="Tuổi từ", required=False, )  # hh_161
    source_man = fields.Integer(string="Nguồn nam", required=False, )  # hh_162
    source_women = fields.Integer(string="Nguồn nữ", required=False, )  # hh_163
    source_total = fields.Integer(string="Số lượng thi tuyển", required=False, )  # hh_164
    age_to = fields.Integer(string="Đến tuổi", required=False, )  # hh_165
    certificate = fields.Many2one(comodel_name="intern.certification", string="Trình độ", required=False, )  # hh_166
    marital = fields.Many2one(comodel_name="marital", string="Hôn nhân", required=False, )  # hh_167
    preferred_hand = fields.Selection(string="Tay thuận",
                                      selection=[('Tay trái', 'Tay trái'), ('Tay phải', 'Tay phải'), ],
                                      required=False, )  # hh_168
    comeback1 = fields.Selection(string="Quay về nước",
                                 selection=[('1 tuần tới', '1 tuần tới'), ('1 tháng tới', '1 tháng tới'),
                                            ('3 tháng tới', '3 tháng tới'), ],
                                 required=False, )
    vision = fields.Char(string="Thị lực", required=False, )  # hh_169
    physical = fields.Char(string="Thể lực", required=False, )  # hh_170
    smoking = fields.Boolean(string="Hút thuốc")  # hh_171
    job_description = fields.Char(string="Nội dung công việc", required=False, )  # hh_172
    other_requirement = fields.Char(string="Tiêu chuẩn khác", required=False, )  # hh_173
    note = fields.Char(string="Ghi chú", required=False, )  # hh_174
    type_recruitment = fields.Char(string="Hình thức tuyển dụng", required=False, )  # hh_175
    count_target = fields.Boolean(string="Tính chỉ tiêu mới nam")  # hh_176
    count_target_women = fields.Boolean(string="Tính chỉ tiêu mới nữ")  # hh_177
    fee_departure = fields.Float(string="Phí xuất cảnh nam(USD)", required=False, )  # hh_178
    fee_departure_women = fields.Float(string="Phí xuất cảnh nữ(USD)", required=False, )  # hh_179
    fee_study = fields.Float(string="Tiền học(USD)", required=False)  # hh_180
    fee_eating = fields.Float(string="Tiền ăn(VND)", required=False)  # hh_181
    bonus_target = fields.Float(string="Thưởng nam", required=False)  # hh_182
    bonus_target_women = fields.Float(string="Thưởng nữ", required=False, )  # hh_183

    legal_name = fields.Selection(
        [('chauhung', 'Châu Hưng'), ('hoanghung', 'Hoàng Hưng'), ('tracodi', 'Tracodi'), ('thuanan', 'Thuận An')],
        'Pháp nhân')

    def domain_interns_clone(self):
        intern_clone = self.env['intern.internclone'].search([('promoted', '=', False)]).ids
        domain = "[('id', 'in', %s)]" % str(intern_clone)
        return domain

    interns_clone = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                     column2="intern_id", string="Thực tập sinh", domain=domain_interns_clone)  # hh_184

    interns_promoted = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                        column2="intern_id", domain=[('promoted', '=', True)])  # hh_185
    interns_confirm_exam = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                            column2="intern_id", domain=[('promoted', '=', True),
                                                                         ('confirm_exam', '=', True),
                                                                         ('issues_raise', '=', False)])  # hh_186
    interns_pass_new = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                        column2="intern_id", domain=[('pass_exam', '=', True),
                                                                     ('preparatory_exam', '=', False),
                                                                     ('cancel_pass', '=', False)])  # hh_187
    interns_escape_exam = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                           column2="intern_id", domain=[('issues_raise', '=', True)])  # hh_191
    interns_preparatory = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                           column2="intern_id", domain=[('preparatory_exam', '=', True),
                                                                        ('pass_exam', '=', False),
                                                                        ('cancel_pass', '=', False)])  # hh_192
    interns_cancel_pass = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                           column2="intern_id", domain=[('cancel_pass', '=', True),
                                                                        ('preparatory_exam', '=', False),
                                                                        ('pass_exam', '=', False)])  # hh_193
    interns_departure = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                         column2="intern_id", domain=[('departure', '=', True)])  # hh_194

    interns_promoted_doc = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                            column2="intern_id", domain=[('promoted', '=', True)])
    interns_exam_doc = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                        column2="intern_id",
                                        domain=[('confirm_exam', '=', True), ('issues_raise', '=', False)])
    interns_pass_doc = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                        column2="intern_id",
                                        domain=[('pass_exam', '=', True), ('cancel_pass', '=', False)])
    interns_pass_doc_hs = fields.Many2many('intern.internclone', relation="intern_invoice_rel", column1="invoice_id",
                                           column2="intern_id",
                                           domain=[('pass_exam', '=', True), ('cancel_pass', '=', False)])

    status = fields.Selection(
        [(4, 'Khởi tạo'), (5, 'Tiến cử'), (1, 'Thi tuyển'), (2, 'Chốt Trúng tuyển'), (3, 'Hoàn thành'),
         (6, 'Tạm dừng'), (7, 'Huỷ bỏ')], string='Trạng thái đơn hàng', default=4)

    color_notice = fields.Char('Cảnh báo màu')

    @api.one
    def toggle_red(self):
        self.color_notice = '#F02121'

    @api.one
    def toggle_yellow(self):
        self.color_notice = '#E9F021'

    @api.one
    def toggle_green(self):
        self.color_notice = '#39F021'

    @api.one
    def toggle_remove_notice(self):
        self.color_notice = False

    previous_stt = fields.Integer('Trạng thái cũ')

    @api.one
    def revert_destroy(self):
        if self.status == 6 or self.status == 7:
            if self.previous_stt == 0:
                self.write({'status': 4})
            else:
                self.write({'status': self.previous_stt})

    @api.one
    def start_promotion(self):
        ensure_one = False
        arr_ids = []
        for intern in self.interns_clone:
            arr_ids.append((intern.id))
            if intern.promoted:
                ensure_one = True
                # break
            if ensure_one:
                self.write({
                    'status': 5,
                })
            else:
                raise ValidationError(u"Chưa có TTS nào trong danh sách tiến cử")

    @api.one
    def confirm_pass(self):
        ensure_one = False
        if not self.date_join_school:
            raise ValidationError(u"Chưa có thông tin ngày nhập học trúng tuyển")
        if not self.date_pass:
            raise ValidationError(u"Chưa có thông tin ngày trúng tuyển")
        if not self.date_departure:
            raise ValidationError(u"Chưa có thông tin ngày dự kiến xuất cảnh")
        if not self.dispatchcom1:
            raise ValidationError(u"Chưa có thông tin pháp nhân")
        if not self.job_vi or not self.job_jp:
            raise ValidationError(u"Chưa có thông tin ngành nghề")
        for intern in self.interns_clone:
            if intern.pass_exam and not intern.issues_raise:
                ensure_one = True
                break
        if ensure_one:
            for intern in self.interns_clone:
                if intern.pass_exam and not intern.issues_raise:
                    if not intern.enterprise:
                        raise ValidationError(u"Chưa có thông tin xí nghiệp của TTS %s" % intern.name)
                    elif not intern.place_to_work:
                        raise ValidationError(u"Chưa có thông tin địa điểm làm việc của TTS %s" % intern.name)
            self.write({
                'status': 2,
            })
            for intern in self.interns_clone:
                intern.write({
                    'done_exam': True,
                    'sequence_pass': intern.sequence_exam
                })
        else:
            raise ValidationError(u"Chưa có TTS nào trong danh sách trúng tuyển")

    #
    # @api.multi
    # def download_pass_report(self):
    #     if not self.date_pass:
    #         raise ValidationError(u"Chưa có thông tin ngày trúng tuyển")
    #     if not self.room_pttt:
    #         raise ValidationError(u"Chưa có thông tin phòng PTTT")
    #     if not self.date_join_school:
    #         raise ValidationError(u"Chưa có thông tin ngày nhập học TT")
    #     if not self.date_departure:
    #         raise ValidationError(u"Chưa có thông tin ngày xuất cảnh dự kiến")
    #     for intern in self.interns_pass_doc:
    #         if not intern.place_to_work:
    #             raise ValidationError(u"Chưa có thông tin địa điểm làm việc của TTS %s" % intern.name)
    #         elif not intern.enterprise:
    #             raise ValidationError(u"Chưa có thông tin xí nghiệp của TTS %s" % intern.name)
    #
    #     return self.env['report'].get_action(self, 'intern.invoice.pass.xlsx',
    #                                          data=None)
    #
    # @api.multi
    # def download_exam_report(self):
    #     return self.env['report'].get_action(self, 'intern.invoice.exam.xlsx',
    #                                          data=None)
    #
    @api.multi
    def download_promoted_report(self):
        arr = []
        for record in self.interns_clone:
            arr.append((record.id))
        arrs = str(arr).replace('[', '(')
        arrss = str(arrs).replace(']', ')')
        file = self.interns_clone.action_export_excel(arr, arrss, self.id)
        url = '/web/content/overtime.report.store/%s/save_file/DS tiến cử' \
              '' \
              '' \
              '.xls?download=true' % file.id
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'download.file',
            'view_mode': 'form',
            'view_type': 'form',
            'views': [(False, 'form')],
            'target': 'new',
            'context': {'default_link': url}
        }

    @api.multi
    def download_exam_report(self):
        arr = []

        for record in self.interns_confirm_exam:
            arr.append((record.id))
        arrs = str(arr).replace('[', '(')
        arrss = str(arrs).replace(']', ')')
        file = self.interns_confirm_exam.action_export_excel_tt(arr, arrss, self.id)
        url = '/web/content/overtime.report.store/%s/save_file/DS Thi tuyển.xls?download=true' % file.id
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'download.file',
            'view_mode': 'form',
            'view_type': 'form',
            'views': [(False, 'form')],
            'target': 'new',
            'context': {'default_link': url}
        }

    promotion_removed = fields.Char('Promotion removed')

    @api.multi
    def read(self, fields=None, load='_classic_read'):
        result = super(Invoice, self).read(fields, load)
        return result

    @api.multi
    def unlink(self):
        return super(Invoice, self).unlink()

    @api.one
    def confirm_exam(self):
        ensure_one = False
        for intern in self.interns_clone:
            if intern.confirm_exam and not intern.issues_raise:
                ensure_one = True
                break
        if ensure_one:
            self.write({
                'status': 1,
            })
            for intern in self.interns_clone:
                intern.write({
                    'exam': True,
                })
        else:
            raise ValidationError(u"Chưa có TTS nào trong danh sách chốt thi tuyển")

    def pause_invoice(self, reason):
        self.write({
            'previous_stt': self.status,
            'status': 6,
            'reason_pause_cancel': reason,
            'date_pause_cancel_exam': fields.datetime.today()
        })
        for intern in self.interns_clone:
            intern.write({
                'cancel_exam': True,
            })

    def cancel_invoice(self, reason):
        self.write({
            'previous_stt': self.status,
            'status': 7,
            'reason_pause_cancel': reason,
            'date_pause_cancel_exam': fields.datetime.today()
        })
        for intern in self.interns_clone:
            intern.write({
                'cancel_exam': True,
            })

    previous_stt = fields.Integer('Trạng thái cũ')

    @api.one
    def revert_destroy(self):
        if self.status == 6 or self.status == 7:
            if self.previous_stt == 0:
                self.write({'status': 4})
            else:
                self.write({'status': self.previous_stt})

    # -------------------Tạo hồ sơ--------------

    def createHeaderDocNew(self, gender=None):
        docs = self.env['intern.document'].search([('name', '=', "CV_HEAD")], limit=1)
        logo = self.env['intern.document'].search([('name', '=', "Logo")], limit=1)
        streamDoc = BytesIO(codecs.decode(docs[0].attachment, "base64"))
        target_document = None
        tempFile1 = None
        lentmp = len(self.interns_exam_doc)
        if gender != None:
            lentmp = 0
            for intern in self.interns_exam_doc:
                if intern.gender == gender:
                    lentmp += 1

        if lentmp > 20:
            target_document = Document(streamDoc)
            tmp = lentmp / 20
            for i in range(1, tmp + 1):
                document = Document(streamDoc)
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if 'tbl_intern' in cell.text:
                                cell.text = cell.text.replace('tbl_intern', 'tbl_intern%d' % i)
                                break
                for element in document.element.body:
                    target_document.element.body.append(element)

            tempFile1 = NamedTemporaryFile(delete=False)
            target_document.save(tempFile1.name)
            tempFile1.flush()
            tempFile1.close()

        if docs:
            tpl = None
            if target_document is None:
                stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
                tpl = DocxTemplate(stream)
            else:
                tpl = DocxTemplate(tempFile1.name)
            context = {}
            interns = self.interns_exam_doc.sorted(key=lambda x: x.sequence_exam)

            counter_index = 0
            for k in range(0, round(lentmp / 20 + 1)):
                table_interns = []
                for i in range(round(20 * k), round(20 * k + 20)):
                    if i >= lentmp:
                        break
                    if gender == None:
                        intern = interns[i]
                    else:
                        while counter_index < len(interns):
                            if interns[counter_index].gender == gender:
                                intern = interns[counter_index]
                                counter_index += 1
                                break
                            counter_index += 1

                    info = {}
                    info['stt'] = str(i + 1)
                    info['htk'] = intern.name.upper()
                    info['htn'] = intern.name_in_japan
                    if intern.gender == 'nu':
                        info['gt'] = u'女'
                    else:
                        info['gt'] = u'男'
                    info['ns'] = intern_utils.date_time_in_jp(intern.day, intern.month, intern.year)
                    info['t'] = str(intern_utils.get_age_jp(datetime.now(), intern.day, intern.month, intern.year))
                    info['nm'] = intern.blood_group

                    left = 1.5 - (10.0 - intern.vision_left) / 10.0
                    right = 1.5 - (10.0 - intern.vision_right) / 10.0
                    info['tlt'] = "%.1f" % (left)
                    info['tlp'] = "%.1f" % (right)
                    if intern.preferred_hand == '0':
                        info['tt'] = u'右'
                    elif intern.preferred_hand == '1':
                        info['tt'] = u'左'
                    else:
                        info['tt'] = u'両手'
                    info['cc'] = str(intern.height)
                    info['cn'] = str(intern.weight)
                    info['iq'] = "%s%%" % intern.iq_percentage
                    info['ktk'] = None
                    info['hn'] = intern.marital_status.name_in_jp
                    if intern.province:
                        info['pro'] = (intern.province.name).upper()
                    info['bc'] = intern.certification.name_in_jp
                    table_interns.append(info)

                if k == 0:
                    context['tbl_intern'] = table_interns
                else:
                    context['tbl_intern%d' % k] = table_interns
            context['logo'] = InlineImage(tpl, BytesIO(codecs.decode(logo[0].attachment, "base64")), width=Mm(35))
            if self.enterprise_doc:
                context['xn'] = self.enterprise_doc.name_jp_enterprise
            context['nd'] = self.guild.name_in_jp
            context['now'] = intern_utils.date_time_in_jp(datetime.now().day, datetime.now().month, datetime.now().year)

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()

            if tempFile1 is not None:
                os.unlink(tempFile1.name)

            return tempFile
        if tempFile1 is not None:
            os.unlink(tempFile1.name)
        return None

    def createCVDoc(self, docs, intern, index):
        stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
        tpl = DocxTemplate(stream)
        context = {}

        context['ye'] = self.year_expire
        if intern.avatar is not None:
            try:
                streamAvatar = BytesIO(codecs.decode(intern[0].avatar, "base64"))
                tpl.replace_pic('avatar.jpg', streamAvatar)
            except:
                _logger.info("error")

        context['stt'] = str(index + 1)
        context['htk'] = intern.name.upper()

        if intern.gender == 'nu':
            context['gt'] = '女'
        else:
            context['gt'] = '男'
        context['t'] = str(intern_utils.get_age_jp(datetime.now(), intern.day, intern.month, intern.year))

        if intern.marital_status:
            context['hn'] = intern.marital_status.name_in_jp
        if intern.name_in_japan:
            context['htn'] = intern.name_in_japan
        if intern.date_of_birth:
            context['ns'] = intern_utils.date_time_in_jp(intern.day, intern.month, intern.year)
        if intern.address:
            tmpaddress = intern.address
            if ',' in intern.address:
                tmp = intern.address.split(',')
                tmpaddress = tmp[len(tmp) - 1].strip()

            context['dc'] = ((tmpaddress) + " - " + (intern.province.name)).upper()
            context['kc'] = intern.province.getDistanceString()
        if intern.phone_number:
            context['sdt'] = intern.phone_number
        if intern.height:
            context['cc'] = str(intern.height)
        if intern.weight:
            context['cn'] = str(intern.weight)
        if intern.vision_left and intern.vision_right:
            left = 1.5 - (10.0 - intern.vision_left) / 10.0
            right = 1.5 - (10.0 - intern.vision_right) / 10.0
            context['tl'] = "%.1f - %.1f" % (left, right)

        if intern.blindness:
            context['cmm'] = CheckedBox().init(24)
            context['kmm'] = CheckBox().init(24)
        else:
            context['cmm'] = CheckBox().init(24)
            context['kmm'] = CheckedBox().init(24)

        if intern.smoking:
            context['cht'] = CheckedBox().init(24)
            context['kht'] = CheckBox().init(24)
        else:
            context['cht'] = CheckBox().init(24)
            context['kht'] = CheckedBox().init(24)

        if intern.preferred_hand == '0':
            context['ptt'] = CheckedBox().init(24)
            context['ttt'] = CheckBox().init(24)
        elif intern.preferred_hand == '1':
            _logger.info('tay phai')
            context['ptt'] = CheckBox().init(24)
            context['ttt'] = CheckedBox().init(24)
        else:
            context['ptt'] = CheckedBox().init(24)
            context['ttt'] = CheckedBox().init(24)

        if intern.surgery:
            context['chx'] = CheckedBox().init(24)
            context['khx'] = CheckBox().init(24)
            context['hxnd'] = intern.surgery_content
        else:
            context['chx'] = CheckBox().init(24)
            context['khx'] = CheckedBox().init(24)

        if intern.drink_alcohol:
            context['crb'] = CheckedBox().init(24)
            context['krb'] = CheckBox().init(24)
        else:
            context['crb'] = CheckBox().init(24)
            context['krb'] = CheckedBox().init(24)
        if intern.certification:
            if intern.certification.id == 1:
                intern.specialized = u'無し'
            elif intern.certification.id == 2:
                intern.specialized = u'無し'
        if intern.specialized:
            context['chn'] = intern_utils.convert_to_docx_string(intern.specialized)
        if intern.favourite:
            context['st'] = intern_utils.convert_to_docx_string(intern.favourite)
        if intern.strong:
            context['dm'] = intern_utils.convert_to_docx_string(intern.strong)
        if intern.weak:
            context['dy'] = intern_utils.convert_to_docx_string(intern.weak)
        if intern.teammate:
            context['kns'] = u'有'
        else:
            context['kns'] = u'無'

        if intern.cooking:
            context['na'] = u'可'
        else:
            context['na'] = u'不可'

        if intern.diseases:
            context['bt'] = u'有'
        else:
            context['bt'] = u'無'
        if intern.blood_group:
            context['nm'] = intern.blood_group
        if intern.check_kureperin:
            pass
        if intern.iq_percentage:
            if u'%' in intern.iq_percentage:
                context['iq'] = intern.iq_percentage
            else:
                context['iq'] = intern.iq_percentage + u"%"
        if intern.family_income:
            context['tng'] = intern.family_income
        if intern.motivation:
            context['dl'] = intern_utils.convert_to_docx_string(intern.motivation)
        if intern.income_after_three_year:
            context['bn'] = intern.income_after_three_year
        if intern.job_after_return:
            context['vn'] = intern_utils.convert_to_docx_string(intern.job_after_return)
        if intern.prefer_object:
            context['mg'] = intern_utils.convert_to_docx_string(intern.prefer_object)
        if intern.memory:
            context['kn'] = intern_utils.convert_to_docx_string(intern.memory)
        if intern.valuable:
            context['qg'] = intern_utils.convert_to_docx_string(intern.valuable)

        context['hm'] = CheckBox().init(20)
        context['hs'] = CheckBox().init(20)
        context['lb'] = CheckBox().init(20)
        if intern.education_status:
            if intern.education_status == '1':
                context['hs'] = CheckedBox().init(20)
            elif intern.education_status == '2':
                context['hm'] = CheckedBox().init(20)
            elif intern.education_status == '3':
                context['lb'] = CheckedBox().init(20)
            context['edct'] = intern.education_content

        table_education = []
        for education in sorted(intern.educations, key=lambda x: x.sequence):
            info = {}
            if education.month_start:
                info['nbd'] = intern_utils.date_time_in_jp(month=education.month_start, year=education.year_start)
            else:
                info['nbd'] = intern_utils.date_time_in_jp(year=education.year_start)
            if education.month_end:
                info['nkt'] = intern_utils.date_time_in_jp(month=education.month_end, year=education.year_end)
            else:
                info['nkt'] = intern_utils.date_time_in_jp(year=education.year_end)
            info['tt'] = intern_utils.convert_to_docx_string(education.school).upper()
            info['lt'] = education.school_type.name_in_jp
            info['cn'] = intern_utils.convert_to_docx_string(education.specialization)
            info['bc'] = education.certificate.name_in_jp
            if education.graduated:
                info['tn'] = u'卒業'
            else:
                info['tn'] = u'未卒業'

            table_education.append(info)

        context['tbl_educations'] = table_education

        table_employment = []
        if len(intern.employments) == 0:
            table_employment.append({})
            table_employment.append({})
        else:
            for employment in sorted(intern.employments, key=lambda x: x.sequence):
                info = {}
                if employment.month_start:
                    info['nbd'] = intern_utils.date_time_in_jp(month=employment.month_start, year=employment.year_start)
                else:
                    info['nbd'] = intern_utils.date_time_in_jp(year=employment.year_start)
                if employment.month_end:
                    info['nkt'] = intern_utils.date_time_in_jp(month=employment.month_end, year=employment.year_end)
                else:
                    info['nkt'] = intern_utils.date_time_in_jp(year=employment.year_end)
                info['ct'] = intern_utils.convert_to_docx_string(employment.company)
                info['cv'] = intern_utils.convert_to_docx_string(employment.description)
                table_employment.append(info)

        context['tbl_employ'] = table_employment

        table_family = []
        for i, person in enumerate(sorted(intern.familys, key=lambda x: x.sequence)):
            if i < 5:
                context['p%dht' % (i + 1)] = (person.name).upper()
                context['p%dqh' % (i + 1)] = person.relationship
                if person.birth_year > 0:
                    context['p%ddt' % (i + 1)] = str((datetime.now().year) - int(person.birth_year))
                if person.job:
                    context['p%dnn' % (i + 1)] = intern_utils.convert_to_docx_string(person.job)
                if person.live_together:
                    context['p%dsc' % (i + 1)] = Tick()
                else:
                    context['p%dsr' % (i + 1)] = Tick()
            else:
                if person.live_together:
                    job = person.job
                    if not person.job:
                        job = ""
                    table_family.append({'ht': (person.name).upper(), 'qh': person.relationship,
                                         'dt': str((datetime.now().year) - int(person.birth_year)),
                                         'nn': job, 'sc': Tick(), 'sr': ''})
                else:
                    job = person.job
                    if not person.job:
                        job = ""
                    table_family.append({'ht': (person.name).upper(), 'qh': person.relationship,
                                         'dt': str((datetime.now().year) - int(person.birth_year)),
                                         'nn': job, 'sc': '', 'sr': Tick()})
        context['tbl_family'] = table_family

        if intern.family_member:
            context['cnt'] = CheckedBox().init(24)
            context['knt'] = CheckBox().init(24)
            context['nton'] = intern.family_member_in_jp
        else:
            context['cnt'] = CheckBox().init(24)
            context['knt'] = CheckedBox().init(24)

        if intern.family_accept:
            context['gdy'] = CheckedBox().init(24)
            context['gpd'] = CheckBox().init(24)
        else:
            context['gdy'] = CheckBox().init(24)
            context['gpd'] = CheckedBox().init(24)

        tpl.set_fix_table_border()
        tpl.render(context)

        tempFile = NamedTemporaryFile(delete=False)
        tpl.save(tempFile)
        tempFile.flush()
        tempFile.close()
        return tempFile

    @api.multi
    def create_doc_hoso(self):
        return {
            'type': 'ir.actions.act_url',
            'url': '/web/binary/download_document_hoso?model=intern.invoice&id=%s&filename=%s.zip' % (
                str(self.id), self.name),
            'target': 'self', }

    def createHeaderDoc(self):
        docs = self.env['intern.document'].search([('name', '=', "CV_HEAD")], limit=1)
        logo = self.env['intern.document'].search([('name', '=', "Logo")], limit=1)
        streamDoc = BytesIO(codecs.decode(docs[0].attachment, "base64"))
        target_document = None
        tempFile1 = None
        if len(self.interns_exam_doc) > 20:
            target_document = Document(streamDoc)
            tmp = len(self.interns_exam_doc) / 20
            for i in range(1, tmp + 1):
                document = Document(streamDoc)
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if 'tbl_intern' in cell.text:
                                cell.text = cell.text.replace('tbl_intern', 'tbl_intern%d' % i)
                                break
                for element in document.element.body:
                    target_document.element.body.append(element)

            tempFile1 = NamedTemporaryFile(delete=False)
            target_document.save(tempFile1.name)
            tempFile1.flush()
            tempFile1.close()

        if docs:
            tpl = None
            if target_document is None:
                stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
                tpl = DocxTemplate(stream)
            else:
                tpl = DocxTemplate(tempFile1.name)

            context = {}
            lengexa = int(len(self.interns_exam_doc) / 20)
            for k in range(0, int(lengexa + 1)):
                table_interns = []
                for i in range(20 * k, 20 * k + 20):
                    if i >= len(self.interns_exam_doc):
                        break
                    intern = self.interns_exam_doc[i]
                    info = {}
                    info['stt'] = str(i + 1)
                    info['htk'] = intern.name.upper()
                    info['htn'] = intern.name_in_japan
                    if intern.gender == 'nu':
                        info['gt'] = u'女'
                    else:
                        info['gt'] = u'男'
                    info['ns'] = intern_utils.date_time_in_jp(intern.day, intern.month, intern.year)
                    info['t'] = str(intern_utils.get_age_jp(datetime.now(), intern.day, intern.month, intern.year))
                    info['nm'] = intern.blood_group

                    if intern.vision_left:
                        left = 1.5 - (10.0 - intern.vision_left) / 10.0
                        info['tlt'] = "%.1f" % (left)
                    else:
                        info['tlt'] = 'False'
                    if intern.vision_right:
                        right = 1.5 - (10.0 - intern.vision_right) / 10.0
                        info['tlp'] = "%.1f" % (right)
                    else:
                        info['tlp'] = 'False'

                    if intern.preferred_hand == '0':
                        info['tt'] = u'右'
                    elif intern.preferred_hand == '1':
                        info['tt'] = u'左'
                    else:
                        info['tt'] = u'両手'
                    info['cc'] = str(intern.height)
                    info['cn'] = str(intern.weight)
                    info['iq'] = "%s%%" % intern.iq_percentage
                    info['ktk'] = intern.check_kureperin
                    info['hn'] = intern.marital_status.name_in_jp
                    info['pro'] = str(intern.province.name).upper()
                    info['bc'] = intern.certification.name_in_jp
                    table_interns.append(info)

                if k == 0:
                    context['tbl_intern'] = table_interns
                else:
                    context['tbl_intern%d' % k] = table_interns
            context['logo'] = InlineImage(tpl, BytesIO(codecs.decode(logo[0].attachment, "base64")), width=Mm(35))
            if self.enterprise_doc:
                context['xn'] = self.enterprise_doc.name_jp_enterprise
            context['nd'] = self.guild.name_in_jp
            context['now'] = intern_utils.date_time_in_jp(datetime.now().day, datetime.now().month, datetime.now().year)

            tpl.render(context)

            tempFile = NamedTemporaryFile(delete=False)
            tpl.save(tempFile)
            tempFile.flush()
            tempFile.close()

            if tempFile1 is not None:
                os.unlink(tempFile1.name)

            return tempFile
        if tempFile1 is not None:
            os.unlink(tempFile1.name)
        return None

    def create_Doc_1_3(self, infor, dataid):
        docs = self.env['intern.document'].search([('name', '=', "doc1-3")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)

            data = self.env['intern.internclone'].search([('id', '=', dataid)], limit=1)
            context = {}
            # oval 3 la Nữ, 1 la Nam, 2 la ko, 5 la có
            if data.gender == 'nam':
                docdata.remove_shape(u'id="1" name="Oval 1"')

            else:
                docdata.remove_shape(u'id="5" name="Oval 5"')

            if data.marital_status.id is not 2:
                docdata.remove_shape(u'id="2" name="Oval 2"')
            else:
                docdata.remove_shape(u'id="6" name="Oval 6"')
            if data.had_visa_jp:
                docdata.remove_shape(u'id="8" name="Oval 8"')
            else:
                docdata.remove_shape(u'id="7" name="Oval 7"')

            context['a3'] = data.name
            context['a5'] = data.name_in_japan
            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)
            context['a42'] = infor.lienket.date_create_letter_promotion
            context['a7'] = intern_utils.date_time_in_jp(data.day,
                                                         data.month,
                                                         data.year)
            context['a8'] = data.date_of_birth_short

            context['a11'] = str(intern_utils.get_age_jp(infor.lienket.date_create_letter_promotion_short,
                                                         data.day,
                                                         data.month,
                                                         data.year))
            if data.hktt:
                context['a15_1'] = data.hktt
                context['a15'] = intern_utils.convert_to_docx_string(data.hktt)

            context['a21'] = u'%s ~ %s' % (
                intern_utils.date_time_in_vn2(data.last_education_from_month,
                                              data.last_education_from_year),
                intern_utils.date_time_in_vn2(data.last_education_to_month,
                                              data.last_education_to_year))
            context['a21_1'] = u'%s ～ %s' % (
                intern_utils.date_time_in_jp(None, data.last_education_from_month,
                                             data.last_education_from_year),
                intern_utils.date_time_in_jp(None, data.last_education_to_month,
                                             data.last_education_to_year))
            context['a22'] = intern_utils.convert_to_docx_string(
                data.last_school_education_jp)
            if data.last_school_education:
                context['a23'] = intern_utils.convert_to_docx_string(
                    data.last_school_education)
            if data.last_education_from_month2 and data.last_education_from_year2 and \
                    data.last_education_to_month2 and data.last_education_to_year2 \
                    and data.last_school_education_jp2 and data.last_school_education2:
                context['b21'] = u'%s ~ %s' % (
                    intern_utils.date_time_in_vn2(data.last_education_from_month2,
                                                  data.last_education_from_year2),
                    intern_utils.date_time_in_vn2(data.last_education_to_month2,
                                                  data.last_education_to_year2))
                context['b21_1'] = u'%s ～ %s' % (
                    intern_utils.date_time_in_jp(None, data.last_education_from_month2,
                                                 data.last_education_from_year2),
                    intern_utils.date_time_in_jp(None, data.last_education_to_month2,
                                                 data.last_education_to_year2))
                context['b22'] = intern_utils.convert_to_docx_string(
                    data.last_school_education_jp2)
                if data.last_school_education2:
                    context['b23'] = intern_utils.convert_to_docx_string(
                        data.last_school_education2)
            table_jobs = []
            for i in range(0, 4):
                if i == 0:
                    info = {}
                    if data.time_employee and data.job_employee_jp:
                        info['a25'] = u'%s ~ %s' % (
                            intern_utils.date_time_in_vn2(data.time_employee_from_month,
                                                          data.time_employee_from_year),
                            intern_utils.date_time_in_vn2(data.time_employee_to_month,
                                                          data.time_employee_to_year))
                        info['a25_1'] = data.time_employee
                        info['a26'] = intern_utils.convert_to_docx_string(
                            data.job_employee_jp)
                        if data.job_employee_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(
                                data.job_employee_vi)
                        table_jobs.append(info)
                    else:
                        break
                if i == 1:
                    info = {}
                    if data.time_employee2 and data.job_employee2_jp:
                        info['a25'] = u'%s ~ %s' % (
                            intern_utils.date_time_in_vn2(data.time_employee2_from_month,
                                                          data.time_employee2_from_year),
                            intern_utils.date_time_in_vn2(data.time_employee2_to_month,
                                                          data.time_employee2_to_year))
                        info['a25_1'] = data.time_employee2
                        info['a26'] = intern_utils.convert_to_docx_string(
                            data.job_employee2_jp)
                        if data.job_employee2_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(
                                data.job_employee2_vi)
                        table_jobs.append(info)
                    else:
                        break
                if i == 2:
                    info = {}
                    if data.time_employee3 and data.job_employee3_jp:
                        info['a25'] = u'%s ~ %s' % (
                            intern_utils.date_time_in_vn2(data.time_employee3_from_month,
                                                          data.time_employee3_from_year),
                            intern_utils.date_time_in_vn2(data.time_employee3_to_month,
                                                          data.time_employee3_to_year))
                        info['a25_1'] = data.time_employee3
                        info['a26'] = intern_utils.convert_to_docx_string(
                            data.job_employee3_jp)
                        if data.job_employee3_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(
                                data.job_employee3_vi)
                        table_jobs.append(info)
                    else:
                        break
                if i == 3:
                    info = {}
                    if data.time_employee4 and data.job_employee4_jp:
                        info['a25'] = u'%s ~ %s' % (
                            intern_utils.date_time_in_vn2(data.time_employee4_from_month,
                                                          data.time_employee4_from_year),
                            intern_utils.date_time_in_vn2(data.time_employee4_to_month,
                                                          data.time_employee4_to_year))
                        info['a25_1'] = data.time_employee4
                        info['a26'] = intern_utils.convert_to_docx_string(
                            data.job_employee4_jp)
                        if data.job_employee4_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(
                                data.job_employee4_vi)
                        table_jobs.append(info)
                    else:
                        break

                if i == 4:
                    info = {}
                    if data.time_employee5 and data.job_employee5_jp:
                        info['a25'] = u'%s ~ %s' % (intern_utils.date_time_in_vn2
                                                    (data.time_employee5_from_month,
                                                     data.time_employee5_from_year),
                                                    intern_utils.date_time_in_vn2(
                                                        data.time_employee5_to_month,
                                                        data.time_employee5_to_year))
                        info['a25_1'] = data.time_employee5
                        info['a26'] = intern_utils.convert_to_docx_string(
                            data.job_employee5_jp)
                        if data.job_employee5_vi:
                            info['a27'] = intern_utils.convert_to_docx_string(
                                data.job_employee5_vi)
                        table_jobs.append(info)
                    else:
                        break
            info = {}
            info['a25'] = u'%s ~ Hiện nay' % (intern_utils.date_time_in_vn2
                                              (data.time_start_at_pc_from_month,
                                               data.time_start_at_pc_from_year))
            info['a25_1'] = data.time_start_at_pc
            info['a26'] = u"%s (%s)" % (intern_utils.convert_to_docx_string
                                        (data.dispatchcom2.name),
                                        intern_utils.convert_to_docx_string(infor.lienket.job_jp))
            if data.dispatchcom2.name_vn:
                info['a27'] = u"%s (%s)" % (intern_utils.convert_to_docx_string
                                            (data.dispatchcom2.name_vn),
                                            intern_utils.convert_to_docx_string(infor.lienket.job_vi))
            table_jobs.append(info)
            if len(table_jobs) == 1:
                table_jobs.append({})
            context['tbl_jobs'] = table_jobs
            context['a38'] = u'%d年%dヶ月' % (data.time_at_pc_year,
                                           data.time_at_pc_month)
            context['a39'] = u'%d năm %d tháng' % (data.time_at_pc_year,
                                                   data.time_at_pc_month)
            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)
            context['a42'] = intern_utils.date_time_in_vn(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)
            context['a84'] = intern_utils.convert_to_docx_string(infor.lienket.job_jp)
            context['a85'] = intern_utils.convert_to_docx_string(infor.lienket.job_vi)

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_Doc_1_10(self, infor, dataid):
        docs = self.env['intern.document'].search([('name', '=', "doc1-10")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            data = self.env['intern.internclone'].search([('id', '=', dataid)], limit=1)
            context = {}

            context['a3'] = data.name
            context['a59'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_jp)
            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)
            context['a74'] = infor.lienket.enterprise.name_jp_enterprise
            context['a102_1'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name_en)
            if '(' in context['a102_1'] and ')' in context['a102_1']:
                first = context['a102_1'].rfind('(')
                last = context['a102_1'].rfind(')')
                context['a102_1'] = context['a102_1'][:first] + context['a102_1'][last + 1:]

            context['a105'] = infor.lienket.dispatchcom1.director

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_Doc_1_13_1(self, infor):
        docs = self.env['intern.document'].search([('name', '=', "doc1-13-1")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            context['a53'] = intern_utils.convert_to_docx_string(infor.lienket.training_center.name_jp)
            context['a57'] = infor.lienket.training_center.responsive_person
            context['a55'] = intern_utils.date_time_in_jp(infor.lienket.training_center.day_create,
                                                          infor.lienket.training_center.month_create,
                                                          infor.lienket.training_center.year_create)

            context['a56'] = intern_utils.convert_to_vn_phone(infor.lienket.training_center.phone_number)
            context['a125'] = intern_utils.convert_to_docx_string(infor.lienket.training_center.address_en)

            mission = infor.lienket.training_center.mission
            if '\n' in infor.lienket.training_center.mission:
                pre = '<w:p><w:r><w:t>'
                post = '</w:t></w:r></w:p>'
                lineBreak = '<w:br/>'
                test = infor.lienket.training_center.mission.replace('\n', lineBreak)
                mission = pre + test + post
            context['a126'] = intern_utils.convert_to_docx_string(mission)
            context['a127'] = infor.lienket.training_center.number_of_employee

            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_Doc_1_13_2(self, infor):
        docs = self.env['intern.document'].search([('name', '=', "doc1-13-2")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            context['a102'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name_en)
            context['a102_1'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name_en)
            if '(' in context['a102_1']:
                first = context['a102_1'].rfind('(')
                last = context['a102_1'].rfind(')')
                context['a102_1'] = context['a102_1'][:first] + context['a102_1'][last + 1:]
            context['a105'] = infor.lienket.dispatchcom1.director

            context['a113'] = intern_utils.date_time_in_jp(infor.lienket.dispatchcom1.day_create,
                                                           infor.lienket.dispatchcom1.month_create,
                                                           infor.lienket.dispatchcom1.year_create)

            context['a110_1'] = intern_utils.convert_to_vn_phone(infor.lienket.dispatchcom1.phone_number)
            context['a112'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.address_en)

            mission = infor.lienket.dispatchcom1.mission
            if '\n' in mission:
                pre = '<w:p><w:r><w:t>'
                post = '</w:t></w:r></w:p>'
                lineBreak = '<w:br/>'
                test = infor.lienket.dispatchcom1.mission.replace('\n', lineBreak)
                mission = pre + test + post
            context['a128'] = intern_utils.convert_to_docx_string(mission)
            context['a129'] = infor.lienket.dispatchcom1.number_of_employee
            context['a130'] = infor.lienket.dispatchcom1.capital
            context['a131'] = infor.lienket.dispatchcom1.revenue

            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_Doc_1_20(self, infor):
        if self.back_to_pc2:
            docs = self.env['intern.document'].search([('name', '=', "doc1-20")], limit=1)
        else:
            docs = self.env['intern.document'].search([('name', '=', "doc1-20-TH2")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            context['a84'] = infor.lienket.job_predefine.name_jp
            context['a85'] = infor.lienket.job_predefine.name
            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)
            context['a42'] = infor.lienket.date_create_letter_promotion

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_Doc_1_21(self, infor, dataid):
        if self.year_expire == 3:
            docs = self.env['intern.document'].search([('name', '=', "doc1-21")], limit=1)
        else:
            docs = self.env['intern.document'].search([('name', '=', "doc1-21-1")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            data = self.env['intern.internclone'].search([('id', '=', dataid)], limit=1)
            context = {}

            context['a3'] = data.name
            context['a5'] = data.name_in_japan
            context['a102'] = infor.lienket.dispatchcom1.name_en
            context['a104'] = infor.lienket.dispatchcom1.name
            context['a74'] = infor.name_jp_enterprise
            context['a75'] = infor.name_romaji_enterprise
            context['a59'] = infor.lienket.guild.name_in_jp
            context['a60'] = infor.lienket.guild.name_in_en
            context['a49'] = intern_utils.date_time_in_jp(infor.lienket.day_pay_finance2,
                                                          infor.lienket.month_pay_finance2,
                                                          infor.lienket.year_pay_finance2)
            context['a50'] = infor.lienket.date_pay_finance2
            context['a47'] = intern_utils.date_time_in_jp(infor.lienket.day_pay_finance1,
                                                          infor.lienket.month_pay_finance1,
                                                          infor.lienket.year_pay_finance1)
            context['a48'] = infor.lienket.date_pay_finance1
            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)
            context['a42'] = infor.lienket.date_create_letter_promotion
            context['a105'] = infor.lienket.dispatchcom1.director
            context['a106'] = infor.lienket.dispatchcom1.position_director
            context['a123'] = infor.lienket.dispatchcom1.position_director_vi
            context['a105_1'] = infor.lienket.dispatchcom1.position_director
            context['a102_1'] = infor.lienket.dispatchcom1.name_en

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_Doc_1_27(self, infor):
        doc1_17_header = self.env['intern.document'].search([('name', '=', "Doc1-17-HEADER")], limit=1)
        doc1_17_footer = self.env['intern.document'].search([('name', '=', "Doc1-17-FOOTER")], limit=1)
        list_interns = None
        if not infor.lienket.hoso_created:
            list_interns = sorted(infor.lienket.interns_pass_doc, key=lambda x: x.sequence_pass)
        else:
            list_interns = sorted(infor.lienket.interns_pass_doc_hs, key=lambda x: x.sequence_pass)

        if doc1_17_header and doc1_17_footer:
            stream = BytesIO(codecs.decode(doc1_17_header[0].attachment, "base64"))
            docdata = DocxTemplate(stream)

            context = {}
            context['date_promotion'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                                     infor.lienket.month_create_letter_promotion,
                                                                     infor.lienket.year_create_letter_promotion)
            context['name'] = str(list_interns[0].name_without_signal).upper()
            context['name_jp'] = str(list_interns[0].name_in_japan).replace(u'・', ' ')

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_Doc_1_28(self, infor, dataid):
        if self.back_to_pc2:
            docs = self.env['intern.document'].search([('name', '=', "doc1-28")], limit=1)
        else:
            docs = self.env['intern.document'].search([('name', '=', "doc1-28-TH2")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            data = self.env['intern.internclone'].search([('id', '=', dataid)], limit=1)
            context = {}

            context['a3'] = data.name
            context['a5'] = data.name_in_japan.replace(u'・', '  ')
            context['a84'] = intern_utils.convert_to_docx_string(infor.lienket.job_jp)
            context['a82'] = intern_utils.convert_to_docx_string(infor.lienket.name_working_department)
            context['a92'] = intern_utils.convert_to_docx_string(data.dispatchcom2.name)
            context['a94'] = data.dispatchcom2.position_person_sign
            context['a95'] = data.dispatchcom2.director
            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)

            if infor.lienket.back_to_pc2:
                context['a92'] = intern_utils.convert_to_docx_string(data.dispatchcom2.name)

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_Doc_1_29(self, infor):
        docs = self.env['intern.document'].search([('name', '=', "doc1-29")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            table_interns = []
            if not infor.lienket.hoso_created:
                interns_pass = sorted(infor.lienket.interns_pass_doc, key=lambda x: x.sequence_pass)
            else:
                interns_pass = sorted(infor.lienket.interns_pass_doc_hs, key=lambda x: x.sequence_pass)
            counter = 0
            for i, intern in enumerate(interns_pass):
                counter += 1
                info = {}
                info['stt'] = str(counter)
                info['htk'] = infor.lienket.interns_pass_doc.name
                info['xc'] = intern_utils.date_time_in_jp(infor.lienket.day_departure_doc,
                                                          infor.lienket.month_departure_doc,
                                                          infor.lienket.year_departure_doc)

                table_interns.append(info)

            context['a43'] = intern_utils.date_time_in_jp(infor.lienket.day_create_plan_training,
                                                          infor.lienket.month_create_plan_training,
                                                          infor.lienket.year_create_plan_training)
            context['a44'] = intern_utils.date_time_in_jp(infor.lienket.day_start_training,
                                                          infor.lienket.month_start_training,
                                                          infor.lienket.year_start_training)
            context['a45'] = intern_utils.date_time_in_jp(infor.lienket.day_end_training,
                                                          infor.lienket.month_end_training,
                                                          infor.lienket.year_end_training)
            context['a53'] = intern_utils.convert_to_docx_string(infor.lienket.training_center.name_jp)
            context['a54'] = intern_utils.convert_to_docx_string(infor.lienket.training_center.address_en)
            context['a59'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_jp)
            context['a61'] = intern_utils.convert_to_docx_string(infor.lienket.guild.address_in_jp)
            context['a66'] = intern_utils.convert_to_docx_string(infor.lienket.guild.position_of_responsive_jp)
            context['a68'] = infor.lienket.guild.name_of_responsive_jp

            context['tbl_intern'] = table_interns

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_hdtn(self, infor, dataid):
        docs = self.env['intern.document'].search([('name', '=', "DocHDTN")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            data = self.env['intern.internclone'].search([('id', '=', dataid)], limit=1)
            context = {}

            context['a3'] = data.name
            context['a7'] = intern_utils.date_time_in_jp(data.day,
                                                         data.month,
                                                         data.year)

            context['a10'] = intern_utils.date_time_in_jp(data.day_identity,
                                                          data.month_identity,
                                                          data.year_identity)
            context['a15'] = data.hktt
            if data.identity:
                context['a16'] = u'%s省公安' % (data.place_cmnd.name)
                context['a9'] = data.identity
            else:
                context['a9'] = data.identity_2
                context['a16'] = u'警察局局長'

            context['a17'] = data.contact_person
            context['a18'] = intern_utils.convert_to_docx_string(data.contact_address)
            context['a19'] = data.contact_relative.relation_jp
            context['a20'] = data.contact_phone
            context['a59'] = infor.lienket.guild.name_in_jp
            context['a70'] = intern_utils.date_time_in_jp(infor.lienket.guild.day_sign,
                                                          infor.lienket.guild.month_sign,
                                                          infor.lienket.guild.year_sign)

            if infor.lienket.guild.note_subsize_jp:
                if not infor.lienket.guild.subsidize_start_month or infor.lienket.guild.subsidize_start_month == 0:
                    context['a73'] = u'%s' % infor.lienket.guild.note_subsize_jp
                else:
                    context['a73'] = u'%s 円 (%s)' % (
                        str("{:,}".format(infor.lienket.guild.subsidize_start_month)),
                        infor.lienket.guild.note_subsize_jp)
            else:
                context['a73'] = u'%s 円' % str("{:,}".format(infor.lienket.guild.subsidize_start_month))

            context['a74'] = intern_utils.convert_to_docx_string(
                data.enterprise.name_jp_enterprise)
            context['a76'] = intern_utils.convert_to_docx_string(
                data.enterprise.address_jp_enterprise)
            context['a77'] = data.enterprise.phone_number_enterprise
            context['a84'] = intern_utils.convert_to_docx_string(infor.lienket.job_jp)
            context['a86'] = infor.lienket.year_expire

            context['a122'] = intern_utils.date_time_in_jp_missing(infor.lienket.day_sign_proletter,
                                                                   infor.lienket.month_sign_proletter,
                                                                   infor.lienket.year_sign_proletter)
            context['a122_3'] = infor.lienket.year_sign_proletter

            context['a102'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name_en)
            context['a105'] = infor.lienket.dispatchcom1.director
            context['a106'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.position_director)
            context['a112'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.address_en)
            context['a110'] = infor.lienket.dispatchcom1.phone_number
            if infor.lienket.dispatchcom1.fax_number:
                context['a111'] = infor.lienket.dispatchcom1.fax_number

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_hdtv(self, infor, dataid):
        docs = self.env['intern.document'].search([('name', '=', "DocHDTV")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}
            data = self.env['intern.internclone'].search([('id', '=', dataid)], limit=1)

            context['a1'] = data.name.upper()
            context['a6'] = intern_utils.date_time_in_en(data.day, data.month, data.year)

            context['a10_1'] = intern_utils.date_time_in_en(data.day_identity,
                                                            data.month_identity,
                                                            data.year_identity)
            context['a15_1'] = data.hktt
            if data.identity:
                context['a16_1'] = u'Công An %s' % (data.place_cmnd.name)
                context['a9'] = data.identity
            else:
                context['a9'] = data.identity_2
                context['a16_1'] = u'Cục trưởng Cục cảnh sát'

            context['a17_1'] = data.contact_person
            context['a18_1'] = data.contact_address
            context['a19_1'] = data.contact_relative.relation
            context['a20'] = data.contact_phone
            context['a60_2'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_en).replace(
                'KYODO KUMIAI', '').replace('KYOUDOU KUMIAI', '')
            context['a71'] = intern_utils.date_time_in_en(infor.lienket.guild.day_sign,
                                                          infor.lienket.guild.month_sign,
                                                          infor.lienket.guild.year_sign)

            if infor.lienket.guild.note_subsize_vi:
                if not infor.lienket.guild.subsidize_start_month or infor.lienket.guild.subsidize_start_month == 0:
                    context['a73'] = u'%s' % infor.lienket.guild.note_subsize_vi
                else:
                    context['a73'] = u'%s Yên (%s)' % (
                        intern_utils.format_number_in_vn(str(infor.lienket.guild.subsidize_start_month)),
                        infor.lienket.guild.note_subsize_vi)
            else:
                context['a73'] = u'%s Yên' % intern_utils.format_number_in_vn(
                    str(infor.lienket.guild.subsidize_start_month))

            context['a75'] = intern_utils.convert_to_docx_string(
                data.enterprise.name_romaji_enterprise)
            context['a79'] = intern_utils.convert_to_docx_string(
                data.enterprise.address_romoji_enterprise)
            context['a77'] = data.enterprise.phone_number_enterprise
            context['a85'] = intern_utils.convert_to_docx_string(infor.lienket.job_vi)
            context['a86'] = infor.lienket.year_expire

            context['a122'] = intern_utils.date_time_in_en_missing(infor.lienket.day_sign_proletter,
                                                                   infor.lienket.month_sign_proletter,
                                                                   infor.lienket.year_sign_proletter)
            context['a122_3'] = infor.lienket.year_sign_proletter
            context['a104'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name)
            context['a105'] = infor.lienket.dispatchcom1.director
            context['a123'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.position_director_vi)
            context['a124'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.address_vi)
            context['a110'] = infor.lienket.dispatchcom1.phone_number
            if infor.lienket.dispatchcom1.fax_number:
                context['a111'] = infor.lienket.dispatchcom1.fax_number

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_certification_end_train(self, infor):
        docs = self.env['intern.document'].search([('name', '=', "CCDT")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            table_interns = []
            list_interns = infor.lienket.interns_pass_doc
            if infor.lienket.hoso_created:
                list_interns = infor.lienket.interns_pass_doc_hs

            interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
            for i, intern in enumerate(interns_pass):
                info = {}
                info['name'] = intern.name
                info['ns'] = intern_utils.date_time_in_jp(intern.day, intern.month, intern.year)
                table_interns.append(info)

            context['tbl_intern'] = table_interns
            context['a59'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_jp)
            context['a53'] = intern_utils.convert_to_docx_string(infor.lienket.training_center.name_jp)
            context['a54'] = intern_utils.convert_to_docx_string(infor.lienket.training_center.address_en)
            context['a56'] = intern_utils.convert_to_vn_phone(infor.lienket.training_center.phone_number)
            context['a44'] = intern_utils.date_time_in_jp(infor.lienket.day_start_training,
                                                          infor.lienket.month_start_training,
                                                          infor.lienket.year_start_training)
            context['a45'] = intern_utils.date_time_in_jp(infor.lienket.day_end_training,
                                                          infor.lienket.month_end_training,
                                                          infor.lienket.year_end_training)

            context['a46'] = intern_utils.date_time_in_jp(infor.lienket.day_create_plan_training_report_customer,
                                                          infor.lienket.month_create_plan_training_report_customer,
                                                          infor.lienket.year_create_plan_training_report_customer)
            context['a57'] = infor.lienket.training_center.responsive_person

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_danh_sach_lao_dong(self, infor):
        docs = self.env['intern.document'].search([('name', '=', "DANH_SACH_LAO_DONG")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            table_interns = []
            list_interns = infor.lienket.interns_pass_doc
            if infor.lienket.hoso_created:
                list_interns = infor.lienket.interns_pass_doc_hs

            interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
            counter = 0
            for i, intern in enumerate(interns_pass):
                counter += 1
                info = {}
                info['stt'] = counter
                info['name'] = intern.name.upper()
                info['ns'] = intern_utils.date_time_in_en(intern.day, intern.month, intern.year)
                if intern.gender == 'nu':
                    info['gender'] = u'Nữ'
                else:
                    info['gender'] = u'Nam'
                info['job'] = infor.lienket.job_en
                info['departure'] = '%s/%s' % (infor.lienket.month_departure_doc, infor.lienket.year_departure_doc)
                tmps = str(intern.hktt).split(",")
                if len(tmps) >= 3:
                    info['province'] = tmps[len(tmps) - 1].strip()
                    info['district'] = tmps[len(tmps) - 2].strip()
                    info['village'] = tmps[len(tmps) - 3].strip()

                table_interns.append(info)

            context['tbl_intern'] = table_interns

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_check_list(self, infor):
        docs = self.env['intern.document'].search([('name', '=', "Checklist")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            enterprise = self.env['xinghiep.xinghiep'].browse(int(infor.id))
            context['xn'] = enterprise.name_jp_enterprise
            context['nd'] = infor.lienket.guild.name_in_jp
            counter = 0

            list_interns = infor.lienket.interns_pass_doc
            if infor.lienket.hoso_created:
                list_interns = infor.lienket.interns_pass_doc_hs

            for intern in list_interns:
                if intern.enterprise and intern.enterprise.id == int(infor.id):
                    counter += 1
            context['sltts'] = counter

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_48(self, infor):
        docs = self.env['intern.document'].search([('name', '=', "doc4-8")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            table_dd = []
            date_start = datetime.strptime(
                '%s/%s/%s' % (infor.lienket.day_start_training,
                              infor.lienket.month_start_training,
                              infor.lienket.year_start_training), '%d/%m/%Y')
            for i in range(0, 21):
                table_dd.append(date_start.strftime('%d/%m/%Y'))
                date_start = date_start + relativedelta(days=1)
            context['tbd'] = table_dd
            context['a44'] = intern_utils.date_time_in_jp(infor.lienket.day_start_training,
                                                          infor.lienket.month_start_training,
                                                          infor.lienket.year_start_training)
            context['a45'] = intern_utils.date_time_in_jp(infor.lienket.day_end_training,
                                                          infor.lienket.month_end_training,
                                                          infor.lienket.year_end_training)

            table_dd2 = []
            for i in range(0, 11):
                table_dd2.append(date_start.strftime('%d/%m/%Y'))
                date_start = date_start + relativedelta(days=1)
            context['tbd2'] = table_dd2
            context['a59'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_jp)
            context['a68'] = infor.lienket.guild.name_of_responsive_jp
            date_tmp = datetime.strptime('%s/%s/%s' % (infor.lienket.day_end_training,
                                                       infor.lienket.month_end_training,
                                                       infor.lienket.year_end_training), '%d/%m/%Y')

            date_tmp = date_tmp + relativedelta(days=1)
            day_tmp = '%02d' % date_tmp.day
            month_tmp = '%02d' % date_tmp.month
            context['a451'] = intern_utils.date_time_in_jp(day_tmp, month_tmp, '%d' % date_tmp.year)

            table_interns = []
            list_interns = None
            if infor.lienket.interns_pass_doc and len(infor.lienket.interns_pass_doc) > 0:
                list_interns = sorted(infor.lienket.interns_pass_doc, key=lambda x: x.sequence_pass)
            else:
                list_interns = sorted(infor.lienket.interns_pass_doc_hs, key=lambda x: x.sequence_pass)

            date_departure_tmp = datetime.strptime(
                '%s-%s-%s' % (infor.lienket.year_departure_doc,
                              infor.lienket.month_departure_doc,
                              infor.lienket.day_departure_doc), '%Y-%m-%d')
            for i, intern in enumerate(list_interns):
                row = {}
                row['stt'] = '%d' % (i + 1)
                row['name'] = u'%s' % intern.name.upper()
                row['date_entry'] = u'%s' % intern_utils.date_time_in_jp('%02d' % date_departure_tmp.day,
                                                                         '%02d' % date_departure_tmp.month,
                                                                         '%d' % date_departure_tmp.year)
                table_interns.append(row)
            context['tb_interns'] = table_interns

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_list_of_sent_en(self, infor):
        enterprise = self.env['xinghiep.xinghiep'].browse(int(infor.id))
        docs = self.env['intern.document'].search([('name', '=', "list_of_sent")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            context['a60'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_en).upper()
            context['a69'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_of_responsive_romaji)
            context['a61'] = intern_utils.convert_to_docx_string(infor.lienket.guild.address_in_jp)
            context['a62'] = intern_utils.convert_to_docx_string(infor.lienket.guild.address_in_romaji)
            if infor.lienket.guild.license_number:
                context['a120'] = infor.lienket.guild.license_number
            if infor.lienket.guild.fax_number:
                context['a64_65'] = "%s/%s" % (infor.lienket.guild.phone_number, infor.lienket.guild.fax_number)
            else:
                context['a64_65'] = infor.lienket.guild.phone_number

            context['a74'] = intern_utils.convert_to_docx_string(enterprise.name_jp_enterprise)
            context['a75'] = intern_utils.convert_to_docx_string(enterprise.name_romaji_enterprise)
            context['a81'] = enterprise.name_of_responsive_en_enterprise
            context['a79'] = intern_utils.convert_to_docx_string(enterprise.address_jp_enterprise)
            context['a77'] = enterprise.phone_number_enterprise
            context['a86'] = infor.lienket.year_expire

            table_interns = []

            list_interns = infor.lienket.interns_pass_doc
            if infor.lienket.hoso_created:
                list_interns = infor.lienket.interns_pass_doc_hs

            interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != int(infor.id):
                    continue
                counter += 1
                info = {}
                info['stt'] = str(counter)
                info['a3'] = intern.name.upper()
                info['a6'] = intern_utils.date_time_in_en(intern.day, intern.month, intern.year)
                if intern.gender == 'nam':
                    info['a13'] = 'MALE'
                else:
                    info['a13'] = 'FEMALE'
                info['a83'] = intern_utils.convert_to_docx_string(infor.lienket.job_en).upper()
                info['a88'] = infor.lienket.year_departure_doc
                info['a89'] = infor.lienket.month_departure_doc
                table_interns.append(info)
            context['tbl_intern'] = table_interns
            context['a41'] = intern_utils.date_time_in_en(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)

            context['a91'] = infor.lienket.position_person_sign
            context['a90'] = infor.lienket.person_sign_proletter.upper()

            context['a102'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name_en)
            context['a121'] = infor.lienket.dispatchcom1.license_number
            context['a105'] = infor.lienket.dispatchcom1.director.upper()
            context['a108'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.address_en)

            phone_tax = infor.lienket.dispatchcom1.phone_number
            if infor.lienket.dispatchcom1.fax_number and infor.lienket.dispatchcom1.fax_number != infor.lienket.dispatchcom1.phone_number:
                phone_tax = phone_tax + "/" + infor.lienket.dispatchcom1.fax_number
            context['a110_111'] = phone_tax

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_list_of_sent_jp(self, infor):
        enterprise = self.env['xinghiep.xinghiep'].browse(int(infor.id))
        docs = self.env['intern.document'].search([('name', '=', "list_of_sent_jp")], limit=1)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}

            context['a59'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_jp)
            context['a60'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_en).upper()
            context['a68'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_of_responsive_jp)
            context['a62'] = intern_utils.convert_to_docx_string(infor.lienket.guild.address_in_romaji)
            context['a61'] = intern_utils.convert_to_docx_string(infor.lienket.guild.address_in_jp)
            if infor.lienket.guild.license_number:
                context['a120'] = infor.lienket.guild.license_number
            if infor.lienket.guild.fax_number:
                context['a64_65'] = "%s/%s" % (infor.lienket.guild.phone_number, infor.lienket.guild.fax_number)
            else:
                context['a64_65'] = infor.lienket.guild.phone_number

            context['a74'] = intern_utils.convert_to_docx_string(enterprise.name_jp_enterprise)
            context['a75'] = intern_utils.convert_to_docx_string(enterprise.name_romaji_enterprise)
            context['a80'] = intern_utils.convert_to_docx_string(enterprise.name_of_responsive_jp_enterprise)
            context['a76'] = intern_utils.convert_to_docx_string(enterprise.address_jp_enterprise)
            context['a77'] = enterprise.phone_number_enterprise
            context['a86'] = infor.lienket.year_expire

            table_interns = []
            list_interns = infor.lienket.interns_pass_doc
            if infor.lienket.hoso_created:
                list_interns = infor.lienket.interns_pass_doc_hs

            interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
            counter = 0
            for i, intern in enumerate(interns_pass):
                if intern.enterprise.id != int(infor.id):
                    continue
                counter += 1
                info = {}
                info['stt'] = str(counter)
                info['a3'] = intern.name.upper()
                info['a7'] = intern_utils.date_time_in_jp(intern.day, intern.month, intern.year)
                if intern.gender == 'nam':
                    info['a12'] = u'男'
                else:
                    info['a12'] = u'女'
                info['a84'] = intern_utils.convert_to_docx_string(infor.lienket.job_jp)
                info['a88'] = infor.lienket.year_departure_doc
                info['a89'] = infor.lienket.month_departure_doc
                table_interns.append(info)

            context['tbl_intern'] = table_interns
            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)

            context['a91_1'] = infor.lienket.position_person_sign_jp
            context['a90_1'] = infor.lienket.person_sign_proletter.upper()

            context['a102'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name_en)
            context['a121'] = infor.lienket.dispatchcom1.license_number
            context['a105'] = infor.lienket.dispatchcom1.director.upper()
            context['a108'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.address_en)

            phone_tax = infor.lienket.dispatchcom1.phone_number
            if infor.lienket.dispatchcom1.fax_number and infor.lienket.dispatchcom1.fax_number != infor.lienket.dispatchcom1.phone_number:
                phone_tax = phone_tax + "/" + infor.lienket.dispatchcom1.fax_number
            context['a110_111'] = phone_tax

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile

    def create_master(self, infor):
        docs = self.env['intern.document'].search([('name', '=', "DocMaster")], limit=1)

        list_interns = infor.lienket.interns_pass_doc
        if infor.lienket.hoso_created:
            list_interns = infor.lienket.interns_pass_doc_hs

        interns_pass = sorted(list_interns, key=lambda x: x.sequence_pass)
        if docs:
            stream = BytesIO(codecs.decode(docs[0].attachment, "base64"))
            docdata = DocxTemplate(stream)
            context = {}
            intern = None
            for ite in interns_pass:
                if ite.enterprise.id == infor.id:
                    intern = ite
                    break
            counter_intern_base_enterprise = 0
            for intern in interns_pass:
                if ite.enterprise.id == infor.id:
                    counter_intern_base_enterprise += 1
            context['a1'] = intern.name.upper()
            context['a2'] = intern.name.upper()
            context['a3'] = intern.name.upper()
            context['a4'] = str(counter_intern_base_enterprise - 1)
            context['a5'] = intern.name_in_japan.replace(u'・', ' ')
            context['a6'] = intern_utils.date_time_in_en(intern.day, intern.month, intern.year)
            context['a7'] = intern_utils.date_time_in_jp(intern.day, intern.month, intern.year)
            context['a8'] = intern_utils.date_time_in_vn(intern.day, intern.month, intern.year)
            if intern.identity:
                context['a9'] = intern.identity
                context['a16'] = u'%s公安省' % (intern.place_cmnd.name)
            else:
                context['a9'] = intern.identity_2
                context['a16'] = u'警察局局長'
            context['a10'] = intern_utils.date_time_in_jp(intern.day_identity, intern.month_identity,
                                                          intern.year_identity)
            context['a11'] = str(
                intern_utils.get_age_jp(infor.lienket.date_create_letter_promotion_short, intern.day, intern.month,
                                        intern.year))
            if intern.gender == 'nam':
                context['a12'] = u'男'
                context['a13'] = u'MALE'
            else:
                context['a12'] = u'女'
                context['a13'] = u'FEMALE'
            if intern.marital_status.id is not 2:
                context['a14'] = u'無'
            else:
                context['a14'] = u'有'

            if intern.hktt:
                context['a15'] = intern.hktt.upper()

            context['a17'] = intern.contact_person
            context['a18'] = intern.contact_address
            context['a19'] = intern.contact_relative.relation_jp
            context['a20'] = intern.contact_phone
            context['a21'] = intern.last_time_education
            context['a22'] = intern_utils.convert_to_docx_string(intern.last_school_education_jp)
            context['a23'] = intern_utils.convert_to_docx_string(intern.last_school_education)
            context['a24'] = intern.last_time_education
            context['a25'] = intern.time_employee
            context['a26'] = intern_utils.convert_to_docx_string(intern.job_employee_jp)
            context['a27'] = intern_utils.convert_to_docx_string(intern.job_employee_vi)

            if intern.time_employee2 and intern.job_employee2_jp and intern.job_employee2_vi:
                context['a28'] = intern.time_employee2
                context['a29'] = intern_utils.convert_to_docx_string(intern.job_employee2_jp)
                context['a30'] = intern_utils.convert_to_docx_string(intern.job_employee2_vi)

            if intern.time_employee3 and intern.job_employee3_jp and intern.job_employee3_vi:
                context['a31'] = intern.time_employee3
                context['a32'] = intern_utils.convert_to_docx_string(intern.job_employee3_jp)
                context['a33'] = intern_utils.convert_to_docx_string(intern.job_employee3_vi)

            context['a34'] = intern_utils.date_time_in_jp(intern.dispatchcom2.day_create,
                                                          intern.dispatchcom2.month_create,
                                                          intern.dispatchcom2.year_create)
            context['a35'] = intern.time_start_at_pc
            context['a36'] = "%s (%s)" % (intern_utils.convert_to_docx_string(intern.dispatchcom2.name),
                                          intern_utils.convert_to_docx_string(infor.lienket.job_jp))
            context['a37'] = "%s (%s)" % (intern_utils.convert_to_docx_string(intern.dispatchcom2.name_vn),
                                          intern_utils.convert_to_docx_string(infor.lienket.job_vi))
            context['a38'] = u'%d年%dヶ月' % (intern.time_at_pc_year, intern.time_at_pc_month)
            context['a39'] = u'%d năm %d tháng' % (intern.time_at_pc_year, intern.time_at_pc_month)

            context['a40'] = intern_utils.date_time_in_jp(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)
            context['a41'] = intern_utils.date_time_in_en(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)

            context['a42'] = intern_utils.date_time_in_vn(infor.lienket.day_create_letter_promotion,
                                                          infor.lienket.month_create_letter_promotion,
                                                          infor.lienket.year_create_letter_promotion)

            context['a43'] = intern_utils.date_time_in_jp(infor.lienket.day_create_plan_training,
                                                          infor.lienket.month_create_plan_training,
                                                          infor.lienket.year_create_plan_training)
            context['a44'] = intern_utils.date_time_in_jp(infor.lienket.day_start_training,
                                                          infor.lienket.month_start_training,
                                                          infor.lienket.year_start_training)
            context['a45'] = intern_utils.date_time_in_jp(infor.lienket.day_end_training,
                                                          infor.lienket.month_end_training,
                                                          infor.lienket.year_end_training)

            context['a46'] = intern_utils.date_time_in_jp(infor.lienket.day_create_plan_training_report_customer,
                                                          infor.lienket.month_create_plan_training_report_customer,
                                                          infor.lienket.year_create_plan_training_report_customer)
            context['a47'] = intern_utils.date_time_in_jp(infor.lienket.day_pay_finance1,
                                                          infor.lienket.month_pay_finance1,
                                                          infor.lienket.year_pay_finance1)
            context['a48'] = infor.lienket.date_pay_finance1
            context['a49'] = intern_utils.date_time_in_jp(infor.lienket.day_pay_finance2,
                                                          infor.lienket.month_pay_finance2,
                                                          infor.lienket.year_pay_finance2)
            context['a50'] = infor.lienket.date_pay_finance2
            context['a51'] = infor.lienket.length_training
            context['a52'] = str(infor.lienket.hours_training)
            context['a53'] = intern_utils.convert_to_docx_string(infor.lienket.training_center.name_jp)
            context['a55'] = intern_utils.date_time_in_jp(infor.lienket.training_center.day_create,
                                                          infor.lienket.training_center.month_create,
                                                          infor.lienket.training_center.year_create)
            context['a56'] = infor.lienket.training_center.phone_number
            context['a57'] = infor.lienket.training_center.responsive_person
            context['a58'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_acronym)
            context['a59'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_jp)
            context['a60'] = intern_utils.convert_to_docx_string(infor.lienket.guild.name_in_en).upper()
            context['a61'] = intern_utils.convert_to_docx_string(infor.lienket.guild.address_in_jp)
            context['a62'] = intern_utils.convert_to_docx_string(infor.lienket.guild.address_in_romaji)
            context['a63'] = infor.lienket.guild.post_code
            context['a64'] = infor.lienket.guild.phone_number
            if infor.lienket.guild.fax_number:
                context['a65'] = infor.lienket.guild.fax_number
            context['a66'] = intern_utils.convert_to_docx_string(infor.lienket.guild.position_of_responsive_jp)
            context['a67'] = intern_utils.convert_to_docx_string(infor.lienket.guild.position_of_responsive_vi)
            context['a68'] = infor.lienket.guild.name_of_responsive_jp
            context['a69'] = infor.lienket.guild.name_of_responsive_romaji
            context['a70'] = intern_utils.date_time_in_jp(infor.lienket.guild.day_sign,
                                                          infor.lienket.guild.month_sign,
                                                          infor.lienket.guild.year_sign)
            context['a71'] = intern_utils.date_time_in_en(infor.lienket.guild.day_sign,
                                                          infor.lienket.guild.month_sign,
                                                          infor.lienket.guild.year_sign)
            context['a72'] = infor.lienket.guild.fee_training_nd_to_pc
            context['a73'] = str(infor.lienket.guild.subsidize_start_month)
            context['a74'] = intern_utils.convert_to_docx_string(intern.enterprise.name_jp_enterprise)
            context['a75'] = intern_utils.convert_to_docx_string(intern.enterprise.name_romaji_enterprise)
            context['a76'] = intern_utils.convert_to_docx_string(intern.enterprise.address_jp_enterprise)
            context['a77'] = intern.enterprise.phone_number_enterprise
            if intern.enterprise.fax_number_enterprise:
                context['a78'] = intern.enterprise.fax_number_enterprise
            context['a79'] = intern_utils.convert_to_docx_string(intern.enterprise.address_jp_enterprise)
            context['a80'] = intern.enterprise.name_of_responsive_jp_enterprise
            context['a81'] = intern.enterprise.name_of_responsive_en_enterprise
            context['a82'] = intern_utils.convert_to_docx_string(infor.lienket.name_working_department)
            context['a83'] = intern_utils.convert_to_docx_string(infor.lienket.job_en)
            context['a84'] = intern_utils.convert_to_docx_string(infor.lienket.job_jp)
            context['a85'] = intern_utils.convert_to_docx_string(infor.lienket.job_vi)
            context['a86'] = infor.lienket.year_expire
            context['a87'] = intern_utils.date_time_in_jp(infor.lienket.day_departure_doc,
                                                          infor.lienket.month_departure_doc,
                                                          infor.lienket.year_departure_doc)
            context['a88'] = infor.lienket.year_departure_doc
            context['a89'] = infor.lienket.month_departure_doc
            context['a90'] = infor.lienket.person_sign_proletter
            context['a91'] = intern_utils.convert_to_docx_string(infor.lienket.position_person_sign)
            context['a92'] = intern_utils.convert_to_docx_string(intern.dispatchcom2.name)
            if intern.dispatchcom2.address:
                context['a93'] = intern_utils.convert_to_docx_string(intern.dispatchcom2.address)
            context['a95'] = intern.dispatchcom2.director.upper()
            context['a96'] = intern_utils.convert_to_docx_string(intern.dispatchcom2.position_person_sign)
            context['a97'] = intern.dispatchcom2.phone_number
            if intern.dispatchcom2.fax_number:
                context['a98'] = intern.dispatchcom2.fax_number
            context['a99'] = intern_utils.date_time_in_jp(intern.dispatchcom2.day_create,
                                                          intern.dispatchcom2.month_create,
                                                          intern.dispatchcom2.year_create)
            context['a100'] = infor.lienket.employee_pttt

            # Phap nhan
            context['a101'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name_jp)
            context['a102'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name_en)
            context['a104'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.name)
            context['a105'] = infor.lienket.dispatchcom1.director.upper()
            context['a106'] = infor.lienket.dispatchcom1.position_director
            # context['a107'] = self.dispatchcom1.address_jp
            context['a108'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.address_en)
            context['a110'] = infor.lienket.dispatchcom1.phone_number
            if infor.lienket.dispatchcom1.fax_number:
                context['a111'] = infor.lienket.dispatchcom1.fax_number

            context['a112'] = intern_utils.convert_to_docx_string(infor.lienket.dispatchcom1.address_en)
            context['a113'] = intern_utils.date_time_in_jp(infor.lienket.dispatchcom1.day_create,
                                                           infor.lienket.dispatchcom1.month_create,
                                                           infor.lienket.dispatchcom1.year_create)

            counter = 113
            # TTS di cung
            if counter_intern_base_enterprise > 1:
                table_interns = []
                iterate_intern = 0
                for i, itern in enumerate(interns_pass):
                    if itern.enterprise.id != infor.id:
                        continue

                    if iterate_intern > 0:
                        info = {}
                        info['a114'] = itern.name.upper()
                        info['a115'] = itern.name.upper()
                        info['a116'] = itern.name_in_japan.replace(u'・', ' ')
                        info['a117'] = intern_utils.date_time_in_en(itern.day, itern.month, itern.year)
                        info['a118'] = intern_utils.date_time_in_jp(itern.day, itern.month, itern.year)
                        info['a119'] = intern_utils.date_time_in_en(itern.day, itern.month, itern.year)
                        info['a120'] = intern_utils.date_time_in_vn(itern.day, itern.month, itern.year)
                        info['a121'] = str(
                            intern_utils.get_age_jp(infor.lienket.date_create_letter_promotion_short, itern.day,
                                                    itern.month,
                                                    itern.year))

                        if itern.gender == 'nam':
                            info['a122'] = u'男'
                            info['a123'] = u'MALE'
                        else:
                            info['a122'] = u'女'
                            info['a123'] = u'FEMALE'
                        if itern.marital_status.id is not 2:
                            info['a124'] = u'無'
                        else:
                            info['a124'] = u'有'

                        if itern.identity:
                            info['a125'] = itern.identity
                            info['a127'] = u'%s省公安' % (itern.place_cmnd.name)
                        else:
                            info['a15'] = itern.identity_2
                            info['a127'] = u'警察局局長'

                        info['a126'] = intern_utils.date_time_in_jp(itern.day_identity, itern.month_identity,
                                                                    itern.year_identity)

                        if itern.hktt:
                            info['a128'] = itern.hktt.upper()
                        info['a129'] = itern.last_time_education
                        info['a130'] = intern_utils.convert_to_docx_string(itern.last_school_education_jp)
                        info['a131'] = intern_utils.convert_to_docx_string(itern.last_school_education)
                        info['a132'] = itern.last_time_education

                        info['a133'] = itern.time_employee
                        info['a134'] = intern_utils.convert_to_docx_string(itern.job_employee_jp)
                        info['a135'] = intern_utils.convert_to_docx_string(itern.job_employee_vi)

                        if itern.time_employee2 and itern.job_employee2_jp and itern.job_employee2_vi:
                            info['a136'] = itern.time_employee2
                            info['a137'] = intern_utils.convert_to_docx_string(itern.job_employee2_jp)
                            info['a138'] = intern_utils.convert_to_docx_string(itern.job_employee2_vi)

                        if itern.time_employee3 and itern.job_employee3_jp and itern.job_employee3_vi:
                            info['a139'] = itern.time_employee3
                            info['a140'] = intern_utils.convert_to_docx_string(itern.job_employee3_jp)
                            info['a141'] = intern_utils.convert_to_docx_string(itern.job_employee3_vi)

                        # tmp = iterate_intern/10
                        info['a142'] = intern_utils.date_time_in_jp(itern.dispatchcom2.day_create,
                                                                    itern.dispatchcom2.month_create,
                                                                    itern.dispatchcom2.year_create)
                        info['a143'] = itern.time_start_at_pc
                        info['a144'] = u'%s (%s)' % (intern_utils.convert_to_docx_string(itern.dispatchcom2.name),
                                                     intern_utils.convert_to_docx_string(infor.lienket.job_jp))
                        info['a145'] = u'%s (%s)' % (intern_utils.convert_to_docx_string(itern.dispatchcom2.name_vn),
                                                     intern_utils.convert_to_docx_string(infor.lienket.job_vi))
                        info['a146'] = u'%d年%dヶ月' % (itern.time_at_pc_year, itern.time_at_pc_month)
                        info['a147'] = u'%d năm %d tháng' % (itern.time_at_pc_year, itern.time_at_pc_month)
                        info['a148'] = itern.contact_person.upper()
                        info['a149'] = itern.contact_address.upper()
                        info['a150'] = itern.contact_relative.relation_jp
                        info['a151'] = itern.contact_phone
                        # stt
                        info['stt'] = str(iterate_intern + 1)
                        for x in range(1, 39):
                            info['s%d' % x] = str(counter + x)
                        counter = counter + 38
                        table_interns.append(info)
                    iterate_intern += 1

                context['tbl_intern'] = table_interns

            docdata.render(context)
            tempFile = NamedTemporaryFile(delete=False)
            docdata.save(tempFile)
            tempFile.flush()
            tempFile.close()
            return tempFile


class Marital(models.Model):
    _name = 'marital'
    _description = u'Tình trạng hôn nhân'
    _rec_name = 'name_in_vn'

    name_in_vn = fields.Char("Tiếng việt")  # hh_102
    name_in_jp = fields.Char("Tiếng nhật")  # hh_103


class PhieuTraLoi(models.Model):
    _name = 'intern.phieutraloi'
    name = fields.Char('Mã số', required=True)
    total_intern = fields.Integer('Tổng số TTS', default=30, required=True)
    total_intern_men = fields.Integer('Tổng số TTS nam', required=True)
    total_intern_women = fields.Integer('Tổng số TTS nữ', required=True)

    total_current_intern = fields.Integer('Số TTS đã đưa vào DS', compute='_compute_total_current_intern')

    @api.multi
    @api.onchange('total_intern', 'total_intern_women')
    def compute_man(self):
        for record in self:
            record.total_intern_men = record.total_intern - record.total_intern_women

    @api.depends('interns')
    def _compute_total_current_intern(self):
        if self.interns:
            self.total_current_intern = len(self.interns)
        else:
            self.total_current_intern = 0

    interns = fields.One2many('intern.internclone', 'phieutraloi_id')

    @api.model
    def create(self, vals):
        if 'interns' in vals:
            ids = []
            for tmp in vals['interns']:
                ids.append(tmp[0])
            interns = self.env['intern.internclone'].search([('id', 'in', ids)])
            count_man = 0
            count_women = 0
            for intern in interns:
                if intern.gender and intern.gender == 'nam':
                    count_man += 1
                else:
                    count_women += 1

            if count_man > vals['total_intern_men'] or count_women > vals['total_intern_women']:
                raise ValidationError(
                    "Số lượng nam/nữ ko phù hợp, có %d nam và %d nữ trong danh sách" % (count_man, count_women))
        return super(PhieuTraLoi, self).create(vals)

    @api.multi
    def write(self, vals):
        return super(PhieuTraLoi, self).write(vals)

    @api.onchange('interns', 'total_intern_men', 'total_intern_women')
    def _onchage_interns(self):

        count_man = 0
        count_women = 0
        for intern in self.interns:
            if intern.gender and intern.gender == 'nam':
                count_man += 1
            else:
                count_women += 1

        if count_man > self.total_intern_men or count_women > self.total_intern_women:
            raise Warning("Số lượng nam/nữ ko phù hợp, có %d nam và %d nữ trong danh sách" % (count_man, count_women))

    @api.multi
    @api.depends('interns', 'total_intern_men', 'total_intern_women')
    def _compute_len_interns(self):
        for rec in self:
            rec.len_interns_man = 0
            rec.len_interns_women = 0
            for intern in rec.interns:
                if intern.gender and intern.gender == 'nam':
                    rec.len_interns_man += 1
                else:
                    rec.len_interns_women += 1
            if rec.len_interns_man == rec.total_intern_men and rec.len_interns_women == rec.total_intern_women:
                rec.has_full = True
            else:
                rec.has_full = False

    len_interns_man = fields.Integer(compute=_compute_len_interns, store=True)
    len_interns_women = fields.Integer(compute=_compute_len_interns, store=True)
    has_full = fields.Boolean(compute=_compute_len_interns, store=True)
